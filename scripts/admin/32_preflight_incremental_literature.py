from __future__ import annotations

import csv
import hashlib
import json
import re
import sys
import time
from collections import Counter
from datetime import datetime, timezone
from pathlib import Path
from typing import Any


PROJECT_ROOT = Path("/data2/lxj/projects/CervixAgent")
MANIFEST_PATH = (
    PROJECT_ROOT
    / "reports"
    / "literature_ingest"
    / "manual_literature_merge_20260726_manifest.json"
)
OUTPUT_BASE = (
    PROJECT_ROOT / "data" / "processed" / "literature" / "manual_merge_20260726"
)


def sha256_file(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as stream:
        for block in iter(lambda: stream.read(1024 * 1024), b""):
            digest.update(block)
    return digest.hexdigest()


def compact(text: str) -> str:
    return re.sub(r"\s+", " ", text).strip()


def pdf_text(source: Path) -> tuple[str, dict[str, Any], list[str]]:
    warnings: list[str] = []
    try:
        from pypdf import PdfReader

        reader = PdfReader(source)
        pages = []
        for page in reader.pages:
            pages.append(page.extract_text(extraction_mode="layout") or "")
        text = "\n\n".join(pages)
        return text, {"page_count": len(reader.pages), "parser": "pypdf"}, warnings
    except Exception as pypdf_error:
        warnings.append(f"pypdf_failed:{type(pypdf_error).__name__}")
        try:
            import fitz

            document = fitz.open(source)
            text = "\n\n".join(page.get_text("text") for page in document)
            page_count = document.page_count
            document.close()
            return text, {"page_count": page_count, "parser": "pymupdf_fallback"}, warnings
        except Exception as pymupdf_error:
            raise RuntimeError(
                f"PDF extraction failed: pypdf={pypdf_error!r}; "
                f"pymupdf={pymupdf_error!r}"
            ) from pymupdf_error


def docx_text(source: Path) -> tuple[str, dict[str, Any], list[str]]:
    from docx import Document

    document = Document(source)
    paragraphs = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]
    table_lines = []
    for index, table in enumerate(document.tables, start=1):
        rows = [" | ".join(cell.text.strip() for cell in row.cells) for row in table.rows]
        table_lines.append(f"[Table {index}]\n" + "\n".join(rows))
    return (
        "\n\n".join(paragraphs + table_lines),
        {"paragraph_count": len(paragraphs), "table_count": len(document.tables), "parser": "python-docx"},
        [],
    )


def xlsx_text(source: Path) -> tuple[str, dict[str, Any], list[str]]:
    from openpyxl import load_workbook

    workbook = load_workbook(source, read_only=True, data_only=False)
    sections = []
    rows_total = 0
    for worksheet in workbook.worksheets:
        rows = []
        for row in worksheet.iter_rows(values_only=True):
            values = ["" if value is None else str(value).strip() for value in row]
            if any(values):
                rows.append(" | ".join(values))
        rows_total += len(rows)
        sections.append(f"[Worksheet: {worksheet.title}]\n" + "\n".join(rows))
    workbook.close()
    return (
        "\n\n".join(sections),
        {"worksheet_count": len(sections), "nonempty_row_count": rows_total, "parser": "openpyxl"},
        [],
    )


def document_title(text: str, fallback: str) -> str:
    for raw_line in text.splitlines()[:40]:
        line = compact(raw_line).lstrip("#").strip()
        if 20 <= len(line) <= 400 and line.lower() not in {"abstract", "introduction"}:
            return line
    return fallback


def write_json(path: Path, payload: Any) -> None:
    path.write_text(json.dumps(payload, ensure_ascii=False, indent=2) + "\n", encoding="utf-8")


def main() -> int:
    manifest = json.loads(MANIFEST_PATH.read_text(encoding="utf-8"))
    entries = [entry for entry in manifest["entries"] if entry["kind"] == "literature"]
    run_id = "preflight_" + datetime.now(timezone.utc).strftime("%Y%m%dT%H%M%SZ")
    run_root = OUTPUT_BASE / "runs" / run_id
    documents_root = run_root / "documents"
    documents_root.mkdir(parents=True, exist_ok=False)

    ledger: list[dict[str, Any]] = []
    review_rows: list[dict[str, str]] = []
    for index, entry in enumerate(entries, start=1):
        source = PROJECT_ROOT / entry["target"]
        suffix = source.suffix.lower()
        started = time.perf_counter()
        record_id = entry["record"]
        item: dict[str, Any] = {
            "item_id": f"MANUAL-20260726-{index:03d}",
            "record": record_id,
            "category": entry["category"],
            "role": entry["role"],
            "source_relative_path": entry["target"],
            "source_sha256": entry["sha256"],
            "extension": suffix,
            "status": "failed",
            "warnings": [],
        }
        try:
            if not source.is_file():
                raise FileNotFoundError(source)
            if sha256_file(source) != entry["sha256"]:
                raise RuntimeError("Source SHA-256 differs from ingest manifest")
            if suffix == ".pdf":
                text, metrics, warnings = pdf_text(source)
            elif suffix == ".docx":
                text, metrics, warnings = docx_text(source)
            elif suffix == ".xlsx":
                text, metrics, warnings = xlsx_text(source)
            elif suffix == ".doc":
                text, metrics, warnings = "", {"parser": "not_supported"}, ["legacy_doc_requires_manual_review"]
                item["status"] = "manual_review_required"
            else:
                text, metrics, warnings = "", {"parser": "not_supported"}, ["unsupported_extension"]
                item["status"] = "manual_review_required"

            text = text.strip()
            item["warnings"].extend(warnings)
            if item["status"] != "manual_review_required":
                if len(text) < 1000:
                    item["warnings"].append("extracted_text_below_1000_characters")
                if entry["role"].startswith("supplement_only_main_missing"):
                    item["warnings"].append("supplement_only_main_missing_in_this_batch")
                title = document_title(text, record_id)
                payload = {
                    "schema_version": 1,
                    "ingest_run": manifest["run"],
                    "item": item,
                    "document": {
                        "title": title,
                        "text": text,
                        "text_preview": compact(text)[:1200],
                    },
                    "metrics": {
                        **metrics,
                        "full_text_characters": len(text),
                        "text_sha256": hashlib.sha256(text.encode("utf-8")).hexdigest(),
                    },
                }
                output = documents_root / f"{item['item_id']}.json"
                write_json(output, payload)
                item["output_relative_path"] = str(output.relative_to(PROJECT_ROOT))
                item["output_sha256"] = sha256_file(output)
                item["metrics"] = payload["metrics"]
                item["status"] = "parsed"
            else:
                item["metrics"] = metrics
        except Exception as exc:
            item["error"] = repr(exc)
        item["duration_seconds"] = round(time.perf_counter() - started, 3)
        ledger.append(item)
        for warning in item["warnings"]:
            review_rows.append(
                {
                    "item_id": item["item_id"],
                    "record": record_id,
                    "source_relative_path": entry["target"],
                    "role": entry["role"],
                    "status": item["status"],
                    "reason": warning,
                }
            )
        if item["status"] == "failed":
            review_rows.append(
                {
                    "item_id": item["item_id"], "record": record_id,
                    "source_relative_path": entry["target"], "role": entry["role"],
                    "status": "failed", "reason": item.get("error", "unknown_error"),
                }
            )
        print(f"{index}/{len(entries)} {item['item_id']} {suffix} {item['status']}", flush=True)

    status_counts = Counter(item["status"] for item in ledger)
    warning_counts = Counter(warning for item in ledger for warning in item["warnings"])
    summary = {
        "schema_version": 1,
        "run_id": run_id,
        "manifest": str(MANIFEST_PATH.relative_to(PROJECT_ROOT)),
        "selected_count": len(entries),
        "status_counts": dict(status_counts),
        "warning_counts": dict(warning_counts),
        "source_files_modified": False,
        "next_gate": "Only parsed fulltext records passing manual/automated quality review may enter high-quality document conversion and RAG chunking.",
    }
    write_json(run_root / "ledger.json", ledger)
    write_json(run_root / "summary.json", summary)
    with (run_root / "manual_review_queue.csv").open("w", encoding="utf-8", newline="") as stream:
        writer = csv.DictWriter(stream, fieldnames=["item_id", "record", "source_relative_path", "role", "status", "reason"])
        writer.writeheader()
        writer.writerows(review_rows)
    (OUTPUT_BASE / "latest_preflight_run.txt").write_text(run_id + "\n", encoding="ascii")
    print(json.dumps(summary, ensure_ascii=False, indent=2))
    return 0 if status_counts.get("failed", 0) == 0 else 3


if __name__ == "__main__":
    sys.exit(main())
